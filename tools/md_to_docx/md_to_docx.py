from __future__ import annotations

"""
NAME
    md_to_docx.py - Markdown to DOCX converter with TOC support.

SYNOPSIS
    python -m tools.md_to_docx.md_to_docx --input README.md [--output out.docx]

DESCRIPTION
    Converts markdown headings and bullet lists into a DOCX file with a title
    page, table of contents, and optional page breaks for top-level sections.

SIDE EFFECTS
    Reads markdown files and writes DOCX output.
"""

import argparse
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _add_toc_field(paragraph) -> None:
    """
    NAME
        _add_toc_field - Insert a Word TOC field into a paragraph.
    """
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)


def _add_page_number(paragraph) -> None:
    """
    NAME
        _add_page_number - Insert a centered page number field.
    """
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def _parse_title_and_purpose(lines: list[str]) -> tuple[str, str]:
    """
    NAME
        _parse_title_and_purpose - Extract title and Purpose line from markdown.
    """
    title = ""
    purpose = ""
    idx = 0
    while idx < len(lines) and not lines[idx].strip():
        idx += 1
    if idx < len(lines):
        line = lines[idx].strip()
        if line.startswith("**") and line.endswith("**"):
            title = line.strip("*").strip()
            idx += 1
        elif line.startswith("# "):
            title = line[2:].strip()
            idx += 1
    while idx < len(lines) and not lines[idx].strip():
        idx += 1
    if idx < len(lines) and lines[idx].strip().startswith("Purpose:"):
        purpose = lines[idx].strip()
    return title, purpose


def build_docx(input_path: Path, output_path: Path, title: str | None) -> None:
    """
    NAME
        build_docx - Convert markdown content into a DOCX document.

    PARAMETERS
        input_path: Markdown file path.
        output_path: DOCX file path to write.
        title: Optional title override.

    SIDE EFFECTS
        Writes a DOCX file to disk.
    """
    lines = input_path.read_text(encoding="utf-8").splitlines()
    detected_title, purpose = _parse_title_and_purpose(lines)
    if not title:
        title = detected_title or input_path.stem

    doc = Document()
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    footer = section.footer
    footer.paragraphs[0].text = ""
    _add_page_number(footer.paragraphs[0])

    # Title page
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(28)
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Source: {input_path.name}")
    if purpose:
        doc.add_paragraph(purpose)
    doc.add_page_break()

    # TOC page
    p = doc.add_paragraph("Table of Contents")
    if p.runs:
        p.runs[0].bold = True
        p.runs[0].font.size = Pt(16)
    toc_p = doc.add_paragraph()
    _add_toc_field(toc_p)
    doc.add_page_break()

    in_code = False
    first_h1 = True
    for line in lines:
        if line.startswith("```"):
            in_code = not in_code
            continue
        if in_code:
            doc.add_paragraph(line)
            continue
        if line.startswith("**") and line.endswith("**") and line.strip("*").strip() == title:
            continue
        if purpose and line.strip() == purpose:
            continue
        if line.startswith("# "):
            if not first_h1:
                doc.add_page_break()
            first_h1 = False
            doc.add_heading(line[2:].strip(), level=1)
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            continue
        if line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=4)
            continue
        if line.startswith("- "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
            continue
        if not line.strip():
            doc.add_paragraph("")
            continue
        doc.add_paragraph(line)

    doc.save(output_path)


def main() -> int:
    """
    NAME
        main - CLI entry point for markdown conversion.
    """
    parser = argparse.ArgumentParser(description="Convert markdown to docx with title, TOC, and page breaks.")
    parser.add_argument("--input", required=True, help="Input markdown file.")
    parser.add_argument("--output", default="", help="Output docx file (default: input name).")
    parser.add_argument("--title", default="", help="Override document title.")
    args = parser.parse_args()

    input_path = Path(args.input)
    output_path = Path(args.output) if args.output else input_path.with_suffix(".docx")
    build_docx(input_path, output_path, args.title or None)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
